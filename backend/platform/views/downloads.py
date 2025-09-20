from django.urls import path
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.http import JsonResponse, FileResponse
from ..permissions import RequiresPremium
from ..services import db
from ..services import constants as backend_constants
import pathlib
import io

# Fallback for transcript DOCX if backend.transcript not present
try:
    from backend import transcript as real_transcript
    def _make_docx(title, url, captions):
        return real_transcript.make_downloadable_text(title, url, captions)
except Exception:
    from docx import Document
    def _make_docx(title, url, captions):
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(url)
        for c in captions:
            text = getattr(c, "text", str(c))
            doc.add_paragraph(text)
        return doc

@api_view(["GET"])
@permission_classes([IsAuthenticated, RequiresPremium])
def download_subtitle_csv(request, video_id:int):
    with db.connect_context() as conn:
        video = db.Video.get(conn, video_id)
    if video.platform == 'youtube':
        path = next((backend_constants.ASSET_ROOT / "subtitles" / "fix").glob(f"{video.on_platform_id}*"), None)
        if not path:
            return JsonResponse({"error": "Transcript not found"}, status=404)
        return FileResponse(open(path, 'rb'), as_attachment=True, filename=path.name)
    return JsonResponse({"error": "Not supported"}, status=400)

@api_view(["GET"])
@permission_classes([IsAuthenticated, RequiresPremium])
def download_subtitle_docx(request, video_id:int):
    import webvtt
    with db.connect_context() as conn:
        video = db.Video.get(conn, video_id)
        try:
            path = next(pathlib.Path("assets/subtitles/raw/").glob(f"{video.on_platform_id}*"))
        except StopIteration:
            return JsonResponse({"error": "Transcript not found"}, status=404)
        captions = webvtt.read(path).captions
        if video.platform == 'youtube':
            docx = _make_docx(video.title, f"https://www.youtube.com/watch?v={video.on_platform_id}", captions)
            output = io.BytesIO()
            docx.save(output)
            output.seek(0)
            return FileResponse(output, as_attachment=True, filename=f"{video.title}.docx")
        return JsonResponse({"error": "Not supported"}, status=400)

@api_view(["GET"])
@permission_classes([IsAuthenticated, RequiresPremium])
def download_audio_mp3(request, video_id:int):
    with db.connect_context() as conn:
        video = db.Video.get(conn, video_id)
    if video.platform == 'youtube':
        path = next((backend_constants.ASSET_ROOT / "audio").glob(f"{video.id}*.mp3"), None)
        if path is None:
            return JsonResponse({"error": "Audio not found"}, status=404)
        return FileResponse(open(path, 'rb'), as_attachment=True, filename=path.name)
    return JsonResponse({"error": "Not supported"}, status=400)

urlpatterns = [
    path("download/subtitle/csv/<int:video_id>", download_subtitle_csv),
    path("download/subtitle/docx/<int:video_id>", download_subtitle_docx),
    path("download/audio/mp3/<int:video_id>", download_audio_mp3),
]
