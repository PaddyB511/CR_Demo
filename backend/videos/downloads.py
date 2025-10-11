from pathlib import Path
import io, tempfile, os
from django.http import FileResponse
from rest_framework import permissions, decorators, response, status
from rest_framework.viewsets import ViewSet
from docx import Document
from .models import Video
from django.utils.text import slugify

ASSET_ROOT = Path("assets")  # optional persistent cache root
SUB_RAW = ASSET_ROOT / "subtitles" / "raw"
AUDIO_DIR = ASSET_ROOT / "audio"
SUB_RAW.mkdir(parents=True, exist_ok=True)
AUDIO_DIR.mkdir(parents=True, exist_ok=True)

YT_THUMB_URL = "https://img.youtube.com/vi/{id}/hqdefault.jpg"

class DownloadError(Exception):
    pass

def _fetch_youtube_captions(video_id: str):
    """Return a list of (text) caption lines for a YouTube video id using webvtt if available via yt_dlp.
    We avoid storing large files unless we decide to cache them. If already cached as .vtt use it.
    """
    import webvtt, json, subprocess, shlex
    # Try cached first
    existing = list(SUB_RAW.glob(f"{video_id}*.vtt"))
    if existing:
        try:
            return [getattr(c, 'text', str(c)) for c in webvtt.read(existing[0]).captions]
        except Exception:
            pass
    # Use yt_dlp to fetch auto subs (requires yt_dlp installed in environment)
    cmd = f"yt-dlp --skip-download --write-auto-subs --sub-lang en --convert-subs vtt -o '%(id)s.%(ext)s' https://www.youtube.com/watch?v={video_id}"
    try:
        subprocess.run(shlex.split(cmd), cwd=str(SUB_RAW), check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
    except Exception as e:
        raise DownloadError(f"Failed to fetch captions: {e}")
    new_files = list(SUB_RAW.glob(f"{video_id}*.vtt"))
    if not new_files:
        raise DownloadError("No subtitle file produced")
    return [getattr(c, 'text', str(c)) for c in webvtt.read(new_files[0]).captions]

def _fetch_youtube_audio(video_id: str, target_path: Path):
    """Download audio as mp3 on demand using yt_dlp; cache result in assets/audio."""
    import subprocess, shlex
    tmpdir = tempfile.mkdtemp()
    # bestaudio -> mp3 conversion
    out_tpl = os.path.join(tmpdir, "%(id)s.%(ext)s")
    cmd = f"yt-dlp -x --audio-format mp3 -o '{out_tpl}' https://www.youtube.com/watch?v={video_id}"
    try:
        subprocess.run(shlex.split(cmd), check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=300)
    except Exception as e:
        raise DownloadError(f"Failed to fetch audio: {e}")
    produced = list(Path(tmpdir).glob(f"{video_id}*.mp3"))
    if not produced:
        raise DownloadError("No audio file produced")
    # Move to cache path
    target_path.parent.mkdir(parents=True, exist_ok=True)
    os.replace(str(produced[0]), target_path)
    return target_path

class IsPremium(permissions.BasePermission):
    def has_permission(self, request, view):
        return bool(request.user and request.user.is_authenticated and getattr(request.user, 'premium', False))


class VideoDownloadViewSet(ViewSet):
    permission_classes = [permissions.IsAuthenticated, IsPremium]

    def _get_video(self, pk):
        return Video.objects.filter(pk=pk).first()

    @decorators.action(detail=True, methods=["get"], url_path="subtitle/docx")
    def subtitle_docx(self, request, pk=None):
        video = self._get_video(pk)
        if not video:
            return response.Response({"error": "Not found"}, status=404)
        if video.platform != 'youtube':
            return response.Response({"error": "Not supported"}, status=400)
        try:
            captions = _fetch_youtube_captions(video.on_platform_id)
        except DownloadError as e:
            return response.Response({"error": str(e)}, status=502)
        doc = Document()
        doc.add_heading(video.title, 0)
        doc.add_paragraph(f"https://www.youtube.com/watch?v={video.on_platform_id}")
        for line in captions:
            if line:
                doc.add_paragraph(line)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        safe_title = slugify(video.title) or f"video-{video.pk}"
        return FileResponse(output, as_attachment=True, filename=f"{safe_title}.docx")

    @decorators.action(detail=True, methods=["get"], url_path="audio/mp3")
    def audio_mp3(self, request, pk=None):
        video = self._get_video(pk)
        if not video:
            return response.Response({"error": "Not found"}, status=404)
        if video.platform != 'youtube':
            return response.Response({"error": "Not supported"}, status=400)
        cache_name = AUDIO_DIR / f"{video.on_platform_id}.mp3"
        if not cache_name.exists():
            try:
                _fetch_youtube_audio(video.on_platform_id, cache_name)
            except DownloadError as e:
                return response.Response({"error": str(e)}, status=502)
        return FileResponse(open(cache_name, 'rb'), as_attachment=True, filename=cache_name.name)
