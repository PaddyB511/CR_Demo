from datetime import datetime
from pathlib import Path
import io, tempfile, os, subprocess, shlex
from django.db.models import Q, Sum
from django.http import FileResponse
from rest_framework import viewsets, permissions, decorators, response, status
from rest_framework.request import Request
from rest_framework.response import Response

from .models import Video, Channel, Tag, Speaker
from .serializers import (
	VideoSerializer,
	VideoDetailSerializer,
	ChannelSerializer,
	TagSerializer,
	SpeakerSerializer,
)
from backend.journal.models import UserViewLog


class VideoViewSet(viewsets.ModelViewSet):
	queryset = Video.objects.all().select_related("channel").prefetch_related("tags", "speakers")
	permission_classes = [permissions.IsAuthenticated]
	serializer_class = VideoSerializer
	filterset_fields = {
		'level': ['exact', 'in'],
		'channel__name': ['exact', 'in'],
		'channel_id': ['exact', 'in'],
		'tags__name': ['exact', 'in'],
		'speakers__name': ['exact', 'in'],
		'premium': ['exact'],
	}
	search_fields = ['title', 'description']
	ordering_fields = ['upload_date', 'duration', 'id']
	ordering = ['-upload_date', '-id']

	def get_serializer_class(self):
		if self.action == "retrieve":
			return VideoDetailSerializer
		return super().get_serializer_class()

	# Filtering logic migrated from legacy platform.utils.filters + browse view
	def get_queryset(self):
		qs = super().get_queryset()
		# Legacy param mapping for backward compatibility
		params = self.request.query_params
		text = params.get('text') or params.get('q')
		if text:
			qs = qs.filter(Q(title__icontains=text) | Q(description__icontains=text))
		# durations legacy parsing
		durations = params.get('durations')
		if durations and ',' in durations:
			d1, d2 = durations.split(',', 1)
			try:
				if d1:
					v1 = float(d1)
					qs = qs.filter(duration__gte=int(v1) * 60 if v1 > 1000 else int(v1))
				if d2:
					v2 = float(d2)
					qs = qs.filter(duration__lte=int(v2) * 60 if v2 > 1000 else int(v2))
			except ValueError:
				pass
		# hide-watched legacy flag
		if params.get('hide-watched', 'false').lower() == 'true' and self.request.user.is_authenticated:
			viewed_ids = UserViewLog.objects.filter(user=self.request.user).values_list('video_id', flat=True)
			qs = qs.exclude(id__in=viewed_ids)
		# manual sort mapping if "sort" provided (overrides ordering)
		sort = params.get('sort')
		if sort:
			sort = sort.lower()
			if sort == 'new':
				qs = qs.order_by('-upload_date', '-id')
			elif sort == 'old':
				qs = qs.order_by('upload_date', 'id')
			elif sort == 'short':
				qs = qs.order_by('duration')
			elif sort == 'long':
				qs = qs.order_by('-duration')
		return qs.distinct()

	def retrieve(self, request, *args, **kwargs):
		obj = self.get_object()
		# premium gate
		if obj.premium and not getattr(request.user, 'premium', False):
			return Response({"error": "Unauthorized"}, status=status.HTTP_401_UNAUTHORIZED)
		return super().retrieve(request, *args, **kwargs)

	@decorators.action(detail=False, methods=["get"], url_path="statistics", permission_classes=[permissions.IsAuthenticated])
	def statistics(self, request: Request):
		qs = self.get_queryset()
		# durations buckets (60s step like legacy) – compute max
		durations = list(qs.values_list("duration", flat=True))
		if durations:
			d_max = max(durations) + 20
		else:
			d_max = 0
		step = 60
		n_steps = max(1, int((d_max // step) + 1))
		buckets = [0] * n_steps
		for d in durations:
			i = min(n_steps - 1, max(0, d // step))
			buckets[i] += 1
		durations_payload = [
			{"count": buckets[i], "start": i * step, "end": (i + 1) * step}
			for i in range(n_steps)
		]
		level_counts = qs.values("level").order_by().annotate(count=Sum(0) + 1)  # hack: will adjust below
		# Simpler manual count to avoid portability issues
		from collections import Counter
		level_counter = Counter(qs.values_list("level", flat=True))
		channel_counter = Counter(qs.values_list("channel__name", flat=True))
		speaker_counter = Counter(
			Speaker.objects.filter(videos__in=qs).values_list("name", flat=True)
		)
		tag_counter = Counter(
			Tag.objects.filter(videos__in=qs).values_list("name", flat=True)
		)
		return response.Response({
			"total": qs.count(),
			"statistics": {
				"levels": [{"name": k, "count": v} for k, v in level_counter.items()],
				"channels": [{"name": k, "count": v} for k, v in channel_counter.items()],
				"speakers": [{"name": k, "count": v} for k, v in speaker_counter.items()],
				"topics": sorted([{"name": k, "count": v} for k, v in tag_counter.items()], key=lambda x: x["name"]),
				"durations": durations_payload,
			},
		})

	@decorators.action(detail=True, methods=["post"], url_path="mark-as-watched", permission_classes=[permissions.IsAuthenticated])
	def mark_as_watched(self, request: Request, pk=None):
		video = self.get_object()
		UserViewLog.objects.create(
			user=request.user,
			video=video,
			watch_date=datetime.utcnow(),
			watch_time=1,
			video_time_start=max(0, video.duration - 1),
			video_time_end=video.duration,
		)
		return Response({}, status=status.HTTP_201_CREATED)

	@decorators.action(detail=False, methods=["post"], url_path="watchtime", permission_classes=[permissions.IsAuthenticated])
	def watchtime(self, request: Request):
		data = request.data or {}
		vid = data.get("videoId")
		if not vid:
			return Response({"error": "Missing videoId"}, status=400)
		try:
			video = Video.objects.get(pk=vid)
		except Video.DoesNotExist:
			return Response({"error": "Video not found"}, status=404)
		try:
			elapsed = float(data.get("elapsedTime", 0))
			t_end = float(data.get("lastVideoTime", 0))
		except (TypeError, ValueError):
			return Response({"error": "Invalid time payload"}, status=400)
		t_start = t_end - elapsed
		UserViewLog.objects.create(
			user=request.user,
			video=video,
			watch_date=datetime.utcnow(),
			watch_time=int(elapsed),
			video_time_start=float(t_start),
			video_time_end=float(t_end),
		)
		return Response({})

	# ---- Download endpoints (premium-only) ----
	def _assert_premium(self, request):
		if not getattr(request.user, 'premium', False):
			return Response({"error": "Premium required"}, status=status.HTTP_402_PAYMENT_REQUIRED)
		return None

	def _yt_captions(self, yt_id: str):
		import webvtt
		raw_dir = Path('assets') / 'subtitles' / 'raw'
		raw_dir.mkdir(parents=True, exist_ok=True)
		# cached?
		cached = list(raw_dir.glob(f"{yt_id}*.vtt"))
		if not cached:
			cmd = f"yt-dlp --skip-download --write-auto-subs --sub-lang en --convert-subs vtt -o '%(id)s.%(ext)s' https://www.youtube.com/watch?v={yt_id}"
			try:
				subprocess.run(shlex.split(cmd), cwd=str(raw_dir), check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
			except Exception as e:
				raise RuntimeError(f"Failed to fetch captions: {e}")
			cached = list(raw_dir.glob(f"{yt_id}*.vtt"))
		if not cached:
			raise RuntimeError("No subtitles produced")
		return webvtt.read(cached[0]).captions

	@decorators.action(detail=True, methods=["get"], url_path="download/subtitle/docx")
	def download_subtitle_docx(self, request: Request, pk=None):
		prem = self._assert_premium(request)
		if prem: return prem
		from docx import Document
		video = self.get_object()
		if video.platform != 'youtube':
			return Response({"error": "Not supported"}, status=400)
		try:
			captions = self._yt_captions(video.on_platform_id)
		except Exception as e:
			return Response({"error": str(e)}, status=502)
		doc = Document()
		doc.add_heading(video.title, 0)
		doc.add_paragraph(f"https://www.youtube.com/watch?v={video.on_platform_id}")
		for c in captions:
			text = getattr(c, 'text', str(c))
			if text:
				doc.add_paragraph(text)
		output = io.BytesIO()
		doc.save(output)
		output.seek(0)
		return FileResponse(output, as_attachment=True, filename=f"{video.on_platform_id}.docx")

	@decorators.action(detail=True, methods=["get"], url_path="download/audio/mp3")
	def download_audio_mp3(self, request: Request, pk=None):
		prem = self._assert_premium(request)
		if prem: return prem
		video = self.get_object()
		if video.platform != 'youtube':
			return Response({"error": "Not supported"}, status=400)
		audio_dir = Path('assets') / 'audio'
		audio_dir.mkdir(parents=True, exist_ok=True)
		cache = audio_dir / f"{video.on_platform_id}.mp3"
		if not cache.exists():
			out_tpl = str(Path(tempfile.mkdtemp()) / "%(id)s.%(ext)s")
			cmd = f"yt-dlp -x --audio-format mp3 -o '{out_tpl}' https://www.youtube.com/watch?v={video.on_platform_id}"
			try:
				res = subprocess.run(shlex.split(cmd), check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=300)
				produced = list(Path(out_tpl).parent.glob(f"{video.on_platform_id}*.mp3"))
				if not produced:
					return Response({"error": "Audio download failed"}, status=502)
				os.replace(str(produced[0]), cache)
			except Exception as e:
				return Response({"error": f"Failed to fetch audio: {e}"}, status=502)
		return FileResponse(open(cache, 'rb'), as_attachment=True, filename=cache.name)


class ChannelViewSet(viewsets.ReadOnlyModelViewSet):
	queryset = Channel.objects.all()
	serializer_class = ChannelSerializer
	permission_classes = [permissions.IsAuthenticated]


class TagViewSet(viewsets.ReadOnlyModelViewSet):
	queryset = Tag.objects.all()
	serializer_class = TagSerializer
	permission_classes = [permissions.IsAuthenticated]


class SpeakerViewSet(viewsets.ReadOnlyModelViewSet):
	queryset = Speaker.objects.all()
	serializer_class = SpeakerSerializer
	permission_classes = [permissions.IsAuthenticated]

